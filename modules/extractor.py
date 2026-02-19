"""Модуль извлечения текста из PDF/DOCX файлов."""
import logging
from pathlib import Path

import pdfplumber
from docx import Document

from modules.models import FileInfo, ExtractedText

logger = logging.getLogger(__name__)


def extract_text(file_info: FileInfo) -> ExtractedText:
    """
    Извлекает текст из файла.

    Для PDF: pdfplumber (постранично, текст объединяется через \\n\\n).
    Для DOCX: python-docx (параграфы + текст из таблиц).

    Если PDF и текст < 50 символов на страницу в среднем → is_scanned=True.

    При любой ошибке возвращает ExtractedText с extraction_method="failed"
    и пустым текстом, НЕ выбрасывает исключение.
    """
    try:
        if file_info.extension == ".pdf":
            return _extract_from_pdf(file_info.path)
        elif file_info.extension == ".docx":
            return _extract_from_docx(file_info.path)
        else:
            logger.warning("Неподдерживаемый формат: %s", file_info.extension)
            return _failed_result()
    except Exception as e:
        logger.error("Ошибка извлечения текста из %s: %s", file_info.filename, e)
        return _failed_result()


def _extract_from_pdf(file_path: Path) -> ExtractedText:
    """Извлекает текст из PDF через pdfplumber."""
    pages_text: list[str] = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages_text.append(text)

    full_text = "\n\n".join(pages_text)
    page_count = len(pages_text)

    # Определение сканированного PDF: если менее 50 символов на страницу
    avg_chars = len(full_text) / max(page_count, 1)
    is_scanned = avg_chars < 50

    if is_scanned:
        logger.info("PDF определён как сканированный (avg %.0f символов/стр): %s",
                     avg_chars, file_path.name)

    return ExtractedText(
        text=full_text.strip(),
        page_count=page_count,
        is_scanned=is_scanned,
        extraction_method="pdfplumber",
    )


def _extract_from_docx(file_path: Path) -> ExtractedText:
    """Извлекает текст из DOCX через python-docx (параграфы + таблицы)."""
    doc = Document(file_path)
    paragraphs: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

    # Также извлечь текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    return ExtractedText(
        text=full_text.strip(),
        page_count=0,  # DOCX не имеет фиксированных страниц
        is_scanned=False,
        extraction_method="python-docx",
    )


def _failed_result() -> ExtractedText:
    """Возвращает результат-заглушку для ошибочных случаев."""
    return ExtractedText(
        text="",
        page_count=0,
        is_scanned=False,
        extraction_method="failed",
    )
