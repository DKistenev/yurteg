"""Единый движок редлайн-DOCX для ЮрТэг.

Использует word-level difflib.SequenceMatcher + python-docx OxmlElement.
Совместим с Microsoft Word 365 (w:ins/w:delText OOXML structure).
"""
import difflib
import io
import re
from itertools import count
from typing import Iterator

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

__all__ = ["generate_redline_docx"]

_AUTHOR = "ЮрТэг"


def _current_date() -> str:
    """Возвращает текущую UTC дату в ISO формате для track changes."""
    from datetime import datetime, timezone
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _tokenize(text: str) -> list[str]:
    """Разбить текст на токены: слова и пробельные последовательности.

    Пробелы сохраняются как отдельные токены, чтобы track changes не
    склеивал слова при вставке/удалении.
    """
    return re.findall(r"\S+|\s+", text)


def _make_rpr() -> "OxmlElement":  # type: ignore[name-defined]
    """Создать пустой w:rPr элемент для track changes."""
    return OxmlElement("w:rPr")


def _add_del_run(para_p: "OxmlElement", text: str, rev_id: str) -> None:  # type: ignore[name-defined]
    """Добавить w:del блок с w:delText в параграф."""
    del_el = OxmlElement("w:del")
    del_el.set(qn("w:id"), rev_id)
    del_el.set(qn("w:author"), _AUTHOR)
    del_el.set(qn("w:date"), _current_date())

    run_el = OxmlElement("w:r")
    rpr = _make_rpr()
    run_el.append(rpr)

    del_text_el = OxmlElement("w:delText")
    del_text_el.set(qn("xml:space"), "preserve")
    del_text_el.text = text
    run_el.append(del_text_el)

    del_el.append(run_el)
    para_p.append(del_el)


def _add_ins_run(para_p: "OxmlElement", text: str, rev_id: str) -> None:  # type: ignore[name-defined]
    """Добавить w:ins блок с w:t в параграф."""
    ins_el = OxmlElement("w:ins")
    ins_el.set(qn("w:id"), rev_id)
    ins_el.set(qn("w:author"), _AUTHOR)
    ins_el.set(qn("w:date"), _current_date())

    run_el = OxmlElement("w:r")
    rpr = _make_rpr()
    run_el.append(rpr)

    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    run_el.append(t_el)

    ins_el.append(run_el)
    para_p.append(ins_el)


def _add_plain_run(para_p: "OxmlElement", text: str) -> None:  # type: ignore[name-defined]
    """Добавить обычный run с текстом."""
    run_el = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    run_el.append(t_el)
    para_p.append(run_el)


def _process_paragraph(
    doc: Document,
    text_old: str,
    text_new: str,
    id_gen: Iterator[int],
) -> None:
    """Обработать один абзац: word-level diff и запись в doc."""
    tokens_old = _tokenize(text_old)
    tokens_new = _tokenize(text_new)

    matcher = difflib.SequenceMatcher(None, tokens_old, tokens_new, autojunk=False)
    opcodes = matcher.get_opcodes()

    # Проверяем — есть ли вообще изменения
    has_changes = any(tag != "equal" for tag, *_ in opcodes)

    para = doc.add_paragraph()
    para_p = para._p

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            text = "".join(tokens_old[i1:i2])
            _add_plain_run(para_p, text)
        elif tag == "delete":
            text = "".join(tokens_old[i1:i2])
            _add_del_run(para_p, text, str(next(id_gen)))
        elif tag == "insert":
            text = "".join(tokens_new[j1:j2])
            _add_ins_run(para_p, text, str(next(id_gen)))
        elif tag == "replace":
            old_text = "".join(tokens_old[i1:i2])
            new_text = "".join(tokens_new[j1:j2])
            _add_del_run(para_p, old_text, str(next(id_gen)))
            _add_ins_run(para_p, new_text, str(next(id_gen)))

    # Если нет изменений — параграф уже содержит plain runs, всё ок
    _ = has_changes  # используется неявно через добавление runs


def generate_redline_docx(
    text_old: str, text_new: str, title: str = "Редлайн"
) -> bytes:
    """Генерирует .docx с track changes (w:ins/w:del) для пары текстов.

    Word-level diff: изменение одного слова затрагивает только это слово.
    Удалённый контент использует w:delText (критично для Word 365).

    Args:
        text_old: Исходный текст.
        text_new: Новый текст.
        title: Заголовок документа.

    Returns:
        bytes готового .docx файла.
    """
    doc = Document()
    doc.add_heading(title, level=1)

    id_gen = count(1)

    # Разбить по абзацам, обработать каждый отдельно
    paragraphs_old = text_old.split("\n")
    paragraphs_new = text_new.split("\n")

    # Выравниваем количество абзацев для попарной обработки
    max_len = max(len(paragraphs_old), len(paragraphs_new))
    paragraphs_old_padded = paragraphs_old + [""] * (max_len - len(paragraphs_old))
    paragraphs_new_padded = paragraphs_new + [""] * (max_len - len(paragraphs_new))

    for p_old, p_new in zip(paragraphs_old_padded, paragraphs_new_padded):
        _process_paragraph(doc, p_old, p_new, id_gen)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
